"""Document parsers — PDF, DOCX, XLSX, TXT + factory.

Mỗi parser trả `ParsedDocument` gồm list `ParsedPage`. Xem
`docs/RAG_PIPELINE.md` §1.2.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import ClassVar, Protocol, runtime_checkable

from openpyxl import load_workbook
from pypdf import PdfReader

from app.core.exceptions import ValidationError


@dataclass
class ParsedPage:
    """Một page/sheet section của document đã parse."""

    page_number: int
    text: str
    section_title: str | None = None


@dataclass
class ParsedDocument:
    """Kết quả parse — input cho chunker."""

    pages: list[ParsedPage] = field(default_factory=list)


@runtime_checkable
class DocumentParser(Protocol):
    """Protocol cho document parser."""

    supported_mime: ClassVar[set[str]]

    def parse(self, file_path: Path) -> ParsedDocument: ...


class TxtParser:
    """Plain text / markdown parser."""

    supported_mime: ClassVar[set[str]] = {"text/plain", "text/markdown"}

    def parse(self, file_path: Path) -> ParsedDocument:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(pages=[ParsedPage(page_number=1, text=text)])


class PDFParser:
    """PDF text extraction via pypdf."""

    supported_mime: ClassVar[set[str]] = {"application/pdf"}

    def parse(self, file_path: Path) -> ParsedDocument:
        reader = PdfReader(str(file_path))
        pages: list[ParsedPage] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages.append(ParsedPage(page_number=idx, text=text))
        return ParsedDocument(pages=pages)


class DocxParser:
    """DOCX parser — dùng heading styles + paragraph text."""

    supported_mime: ClassVar[set[str]] = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    _HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)

    def parse(self, file_path: Path) -> ParsedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        lines: list[str] = []
        current_title: str | None = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style is not None else ""
            heading_match = self._HEADING_RE.match(style_name)
            if heading_match:
                level = int(heading_match.group(1))
                prefix = "#" * level
                lines.append(f"{prefix} {text}")
                current_title = text
            else:
                lines.append(text)

        body = "\n\n".join(lines)
        return ParsedDocument(
            pages=[ParsedPage(page_number=1, text=body, section_title=current_title)],
        )


class XlsxParser:
    """XLSX parser — mỗi row là một unit với sheet context."""

    supported_mime: ClassVar[set[str]] = {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }

    def parse(self, file_path: Path) -> ParsedDocument:
        wb = load_workbook(str(file_path), read_only=True, data_only=True)
        pages: list[ParsedPage] = []
        page_num = 1
        for sheet in wb.worksheets:
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                cells = [
                    str(c).strip() for c in row if c is not None and str(c).strip()
                ]
                if not cells:
                    continue
                text = " | ".join(cells)
                pages.append(
                    ParsedPage(
                        page_number=page_num,
                        text=f"Sheet: {sheet.title}\nRow: {row_idx}\n{text}",
                        section_title=f"Sheet: {sheet.title}",
                    ),
                )
                page_num += 1
        wb.close()
        return ParsedDocument(pages=pages)


_PARSER_REGISTRY: list[type[DocumentParser]] = [
    PDFParser,
    DocxParser,
    XlsxParser,
    TxtParser,
]


def get_parser(mime_type: str) -> DocumentParser:
    """Factory — trả parser phù hợp với MIME type.

    Raises:
        ValidationError: nếu MIME không được hỗ trợ.
    """
    normalized = mime_type.split(";", maxsplit=1)[0].strip().lower()
    for parser_cls in _PARSER_REGISTRY:
        if normalized in parser_cls.supported_mime:
            return parser_cls()
    raise ValidationError(
        f"Unsupported MIME type: {mime_type}",
        details={"mime_type": mime_type},
    )


def supported_mime_types() -> set[str]:
    """Union of all supported MIME types."""
    result: set[str] = set()
    for parser_cls in _PARSER_REGISTRY:
        result.update(parser_cls.supported_mime)
    return result
